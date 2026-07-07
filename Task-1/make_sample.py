"""
Generates a dummy RFP proposal document for testing the portal.
Creates  sample_rfp/sample_rfp.docx  and  sample_rfp/sample_rfp.txt

Run:  python make_sample.py
"""
import os
import docx

SAMPLE = """REQUEST FOR PROPOSAL (RFP No. 2026-IAM-0471)
Statewide Identity and Access Management and Cybersecurity Services
Issued by: Maryland Department of Digital Services

1. INTRODUCTION
The Maryland Department of Digital Services ("the Department") invites qualified vendors to submit
proposals to design, implement and maintain an enterprise Identity and Access Management (IAM)
and cybersecurity solution. The selected contractor shall provide a complete managed service over
a three-year term with two optional one-year renewals.

2. SCOPE OF WORK AND DELIVERABLES
The contractor shall provide the following deliverables:
- The vendor shall deliver a centralized Identity and Access Management platform with single sign-on (SSO).
- The contractor shall implement multi-factor authentication (MFA) across all state agencies.
- The vendor shall provide encryption of data at rest and in transit using industry standard algorithms.
- The contractor shall develop and deliver a detailed project implementation plan within 30 days of award.
- The vendor shall furnish monthly security monitoring and incident response reports.
- The contractor shall conduct quarterly penetration testing and vulnerability assessments.
- The vendor shall provide integration with existing HR and directory systems via API.
- The contractor shall deliver a role-based access control (RBAC) policy framework.
- The vendor shall implement automated user provisioning and de-provisioning workflows.
- The contractor shall maintain 99.9% system uptime and provide 24x7 support.
- The vendor shall deliver a disaster recovery and business continuity plan.
- The contractor shall provide knowledge-transfer training to Department staff.

3. TECHNICAL REQUIREMENTS
The technical requirements include compliance with NIST 800-53 and ISO 27001 standards. The system
specification must support role-based access control. The solution shall be SOC 2 Type II certified.
Integration with the state single sign-on (SSO) gateway is required. The platform must support
SAML 2.0 and OpenID Connect for federation, and provide a documented REST API for interoperability.

4. EVALUATION CRITERIA
Proposals will be evaluated and scored on the following weighted criteria:
- Technical approach and solution design: 30 points
- Relevant experience and past performance: 20 points
- Cost and price proposal: 20 percent
- Key personnel qualifications: 10 points
- Security and compliance posture: 10 points
- Implementation schedule and methodology: 5 points
- Small business / MBE participation: 5 points
The contract will be awarded to the proposal offering the best overall value to the State.

5. FINANCIAL AND PAYMENT TERMS
Payment terms are NET30 from receipt of a valid invoice. Payment will be made against project
milestones. A retainage of 5 percent applies until final acceptance. The estimated contract value
is $4,200,000 over three years. Vendors must submit audited financial statements for the prior two
years as proof of financial stability. A bid bond of 5 percent of the bid amount is required, and a
performance bond will be required from the awarded vendor.

6. INSURANCE REQUIREMENTS
The selected contractor shall maintain general liability insurance coverage of $5 million per
occurrence, professional liability (indemnity) insurance of $2 million, and cyber liability
insurance of $3 million. Proof of workers compensation insurance is also required.

7. LEGAL AND COMPLIANCE REQUIREMENTS
The vendor must comply with all applicable laws and regulations, including state and federal data
protection laws and the Maryland Personal Information Protection Act. The contractor shall be
registered to do business in the State of Maryland (state registration). The vendor must use the
E-Verify system for all personnel. Eligibility criteria include a minimum of five years of relevant
experience and at least three comparable public-sector references. The terms and conditions include
termination for convenience, liability limits, indemnification, intellectual property ownership, and
dispute resolution by arbitration.

8. SUBMISSION REQUIREMENTS AND REQUIRED FORMS
Proposals are due no later than 3:00 PM EST on August 15, 2026 (submission deadline). A mandatory
pre-proposal conference will be held on July 20, 2026. Proposals must be submitted in PDF format,
not to exceed 50 pages, font size 11 or larger, with one original and one electronic copy.
The following forms and certifications must be completed and signed by an authorized signatory:
- Vendor Information Form (Tax ID, owner name, percentage of ownership)
- Minority Business Enterprise (MBE) certification, specify type
- Workers Compensation Insurance certificate
- Non-collusion affidavit and conflict of interest declaration
- Business with Iran certification
- Bid bond in the amount of 5 percent of the bid
All vendors must complete vendor registration on the Maryland eMarketplace supplier portal prior to award.
"""


def main():
    out_dir = os.path.join(os.path.dirname(__file__), "sample_rfp")
    os.makedirs(out_dir, exist_ok=True)

    # TXT
    txt_path = os.path.join(out_dir, "sample_rfp.txt")
    with open(txt_path, "w", encoding="utf-8") as f:
        f.write(SAMPLE)

    # DOCX
    doc = docx.Document()
    for line in SAMPLE.split("\n"):
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph("")
        elif stripped[0].isdigit() and ". " in stripped[:4]:
            doc.add_heading(stripped, level=1)
        elif stripped.startswith("REQUEST FOR PROPOSAL"):
            doc.add_heading(stripped, level=0)
        else:
            doc.add_paragraph(stripped)
    docx_path = os.path.join(out_dir, "sample_rfp.docx")
    doc.save(docx_path)

    print("Created:")
    print(" ", txt_path)
    print(" ", docx_path)


if __name__ == "__main__":
    main()
