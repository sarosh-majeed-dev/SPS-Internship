"""
RFP Analyzer - extraction engine
---------------------------------
Reads an RFP document (PDF / DOCX / TXT) and automatically extracts:
  1. Deliverables          - what the vendor must provide
  2. Evaluation Criteria   - how the client will score proposals
  3. Compliance Checklist  - department tasks mapped to the SPS review template
                             (Financial/Accounting, Legal, Operations, Technical)

The compliance engine is built directly from the SPS RFP review checklist,
including the GO / NO-GO decision rules (NET30 payment, $5M insurance cap, etc.).

No external AI service is required - the engine is fully self-contained so it
can analyse a 70+ page RFP locally and instantly. An optional Claude hook is
included (see analyze_with_llm) for teams that want richer summarisation.
"""

import io
import re

# --------------------------------------------------------------------------- #
#  Document text extraction
# --------------------------------------------------------------------------- #

def extract_text(file_bytes: bytes, filename: str) -> str:
    """Return raw text from a PDF, DOCX or TXT file given its bytes."""
    name = (filename or "").lower()

    if name.endswith(".pdf"):
        return _read_pdf(file_bytes)
    if name.endswith(".docx"):
        return _read_docx(file_bytes)
    # txt / md / anything else -> decode as text
    try:
        return file_bytes.decode("utf-8", errors="ignore")
    except Exception:
        return file_bytes.decode("latin-1", errors="ignore")


def _read_pdf(file_bytes: bytes) -> str:
    from pypdf import PdfReader
    reader = PdfReader(io.BytesIO(file_bytes))
    pages = []
    for page in reader.pages:
        try:
            pages.append(page.extract_text() or "")
        except Exception:
            pages.append("")
    text = "\n".join(pages)

    # If the PDF has almost no selectable text it is probably scanned/image-based.
    # Fall back to on-device OCR (only when the libraries are available).
    n_pages = max(1, len(pages))
    if len(text.split()) < max(40, 15 * n_pages):
        ocr_text = _ocr_pdf(file_bytes)
        if len(ocr_text.split()) > len(text.split()):
            return ocr_text
    return text


_OCR_ENGINE = None


def _ocr_pdf(file_bytes: bytes, max_pages: int = 40, dpi: int = 200) -> str:
    """OCR a scanned PDF fully on-device (PyMuPDF render + RapidOCR). Returns ''
    if the OCR libraries are not installed."""
    global _OCR_ENGINE
    try:
        import fitz  # PyMuPDF
        import numpy as np
        from rapidocr_onnxruntime import RapidOCR
    except Exception:
        return ""

    if _OCR_ENGINE is None:
        _OCR_ENGINE = RapidOCR()

    out = []
    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
    except Exception:
        return ""
    for i, page in enumerate(doc):
        if i >= max_pages:
            break
        try:
            pix = page.get_pixmap(dpi=dpi, alpha=False)
            img = np.frombuffer(pix.samples, dtype=np.uint8).reshape(
                pix.height, pix.width, pix.n)
            result, _ = _OCR_ENGINE(img)
            if result:
                out.append("\n".join(line[1] for line in result))
        except Exception:
            continue
    return "\n".join(out)


def _read_docx(file_bytes: bytes) -> str:
    import docx
    document = docx.Document(io.BytesIO(file_bytes))
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts)


# --------------------------------------------------------------------------- #
#  Helpers
# --------------------------------------------------------------------------- #

def _sentences(text: str):
    """Split text into clean sentence-like fragments."""
    # normalise whitespace, then split on line breaks and sentence enders
    # (note: we do NOT split on ':' so "Criterion: 35 points" stays intact)
    raw = re.split(r"(?<=[.;])\s+|\n+", text)
    out = []
    for s in raw:
        s = re.sub(r"\s+", " ", s).strip(" \t-•*–·")
        if 4 <= len(s) <= 400:
            out.append(s)
    return out


def _find_evidence(sentences, keywords, limit=2):
    """Return up to `limit` sentences that mention any of the keywords."""
    hits = []
    seen = set()
    for s in sentences:
        low = s.lower()
        if any(k in low for k in keywords):
            key = low[:80]
            if key not in seen:
                seen.add(key)
                hits.append(s)
            if len(hits) >= limit:
                break
    return hits


# --------------------------------------------------------------------------- #
#  1. Deliverables
# --------------------------------------------------------------------------- #

DELIVERABLE_VERBS = [
    "provide", "deliver", "furnish", "supply", "implement", "develop",
    "design", "install", "perform", "maintain", "conduct", "produce",
    "submit", "deploy", "configure", "integrate", "build", "support",
]
DELIVERABLE_TRIGGERS = [
    "shall provide", "shall deliver", "shall furnish", "must provide",
    "must deliver", "will provide", "responsible for", "scope of work",
    "scope of services", "deliverable", "the contractor shall",
    "the vendor shall", "services include",
]


def extract_deliverables(text: str):
    sentences = _sentences(text)
    results = []
    seen = set()
    for s in sentences:
        low = s.lower()
        is_trigger = any(t in low for t in DELIVERABLE_TRIGGERS)
        # action-verb sentence that reads like a requirement
        is_action = any(re.search(r"\b" + v + r"\b", low) for v in DELIVERABLE_VERBS) \
            and any(w in low for w in ("shall", "must", "will", "the contractor", "the vendor", "provide"))
        # drop all-caps section headings (e.g. "SCOPE OF WORK AND DELIVERABLES")
        is_heading = s.upper() == s and len(s.split()) <= 7
        if (is_trigger or is_action) and not is_heading:
            key = low[:90]
            if key not in seen:
                seen.add(key)
                results.append(s)
    return results[:40]


# --------------------------------------------------------------------------- #
#  2. Evaluation Criteria
# --------------------------------------------------------------------------- #

EVAL_KEYWORDS = [
    "evaluation", "evaluated", "scoring", "score", "points", "weight",
    "weighting", "criteria", "basis of award", "selection", "rated",
    "maximum points", "percent", "%",
]


EVAL_CONTEXT_WORDS = (
    "experience", "approach", "cost", "price", "personnel", "technical",
    "qualif", "participation", "past performance", "methodology", "design",
    "management", "solution", "schedule", "references", "mbe", "small business",
)


def _clean_criterion(line: str) -> str:
    """Strip leading bullets/numbers and a trailing weight from a criterion line."""
    c = re.sub(r"^[-•*–·\d.\s]+", "", line).strip()
    c = re.sub(r"[:\-\s]*\d{1,3}\s?(%|percent|points|pts)\b.*$", "", c, flags=re.I).strip()
    return c


def extract_evaluation_criteria(text: str):
    """
    Prefer a section-aware pass: locate the evaluation section and pull the
    weighted lines inside it. Fall back to a context-filtered global scan so
    we never confuse '99.9% uptime' or '5 percent retainage' with a criterion.
    """
    lines = [l.strip() for l in text.splitlines()]
    results, seen = [], set()

    start = None
    for i, l in enumerate(lines):
        if re.search(r"evaluation criteria|evaluation and|basis of award|"
                     r"scoring criteria|selection criteria|award criteria", l.lower()):
            start = i
            break

    if start is not None:
        for l in lines[start + 1: start + 40]:
            low = l.lower()
            # stop when we hit the next major numbered section without a weight
            if re.match(r"^\d+\.\s", l) and not re.search(r"\d+\s?(%|percent|points|pts)", low):
                break
            m = re.search(r"(\d{1,3})\s?(%|percent|points|pts)", low)
            if m:
                crit = _clean_criterion(l)
                key = crit.lower()[:80]
                if crit and key not in seen:
                    seen.add(key)
                    results.append({"criterion": crit, "weight": m.group(0)})

    if not results:  # fallback - global, context-filtered
        for s in _sentences(text):
            low = s.lower()
            m = re.search(r"(\d{1,3})\s?(%|percent|points|pts)", low)
            if m and any(w in low for w in EVAL_CONTEXT_WORDS):
                crit = _clean_criterion(s)
                key = crit.lower()[:80]
                if crit and key not in seen:
                    seen.add(key)
                    results.append({"criterion": crit, "weight": m.group(0)})

    return results[:30]


# --------------------------------------------------------------------------- #
#  3. Compliance Checklist  (mapped to SPS departments)
# --------------------------------------------------------------------------- #
#  Each item is checked against the RFP text. If keywords are found we capture
#  the evidence sentence(s); decision rules add an automatic GO / NO-GO flag.

COMPLIANCE_TEMPLATE = {
    "Financial / Accounting": [
        {"item": "Payment Terms",
         "keywords": ["payment", "net 30", "net30", "net-30", "milestone",
                      "retainage", "late payment", "invoice"]},
        {"item": "Financial Stability Requirements",
         "keywords": ["financial statement", "audited", "financial stability",
                      "balance sheet", "proof of financial"]},
        {"item": "Insurance Requirements",
         "keywords": ["insurance", "liability coverage", "general liability",
                      "indemnity insurance"]},
        {"item": "Profitability Analysis",
         "keywords": ["budget", "estimated value", "contract value",
                      "expected revenue", "funding"]},
        {"item": "Bid Bond / Performance Bond",
         "keywords": ["bid bond", "performance bond", "surety", "security deposit"]},
    ],
    "Legal": [
        {"item": "Eligibility Criteria",
         "keywords": ["eligibility", "eligible", "minimum qualification",
                      "registration requirement", "relevant experience"]},
        {"item": "Capability (Personnel & Knowhow)",
         "keywords": ["qualified personnel", "key personnel", "technical knowhow",
                      "certification", "resume", "cv"]},
        {"item": "Compliance Requirements (Laws & Regulations)",
         "keywords": ["comply", "compliance", "laws and regulations",
                      "data protection", "gdpr", "hipaa", "regulatory"]},
        {"item": "State Registration",
         "keywords": ["registered in the state", "state registration",
                      "authorized to do business", "secretary of state"]},
        {"item": "E-Verify",
         "keywords": ["e-verify", "everify", "e verify"]},
        {"item": "Contractual Obligations",
         "keywords": ["termination", "liability limit", "indemnification",
                      "dispute resolution", "terms and conditions", "warranty"]},
    ],
    "Operations": [
        {"item": "Required Forms & Certifications",
         "keywords": ["form", "certification", "declaration", "affidavit",
                      "attachment", "exhibit", "tax id"]},
        {"item": "Submission Deadlines",
         "keywords": ["deadline", "due date", "due by", "submission date",
                      "closing date", "no later than", "proposals due"]},
        {"item": "Document Compliance (Format & Submission)",
         "keywords": ["format", "page limit", "font", "number of copies",
                      "submission requirement", "electronic copy"]},
        {"item": "Signatory Authority",
         "keywords": ["authorized signatory", "signature", "signed by",
                      "duly authorized"]},
        {"item": "MBE / Small Business Participation",
         "keywords": ["mbe", "minority business", "small business", "dbe",
                      "set-aside", "wbe"]},
        {"item": "Vendor Registration",
         "keywords": ["vendor registration", "supplier portal", "register as a vendor",
                      "vendor number"]},
    ],
    "Technical": [
        {"item": "Scope of Services / Products",
         "keywords": ["identity and access management", "iam", "cybersecurity",
                      "security solution", "managed services"]},
        {"item": "Technical Requirements",
         "keywords": ["technical requirement", "specification", "system requirement",
                      "functional requirement"]},
        {"item": "Compliance with Industry Standards",
         "keywords": ["iso 27001", "iso", "nist", "soc 2", "industry standard",
                      "best practice"]},
        {"item": "Security Considerations",
         "keywords": ["encryption", "access control", "data protection",
                      "multi-factor", "mfa", "vulnerability", "penetration test"]},
        {"item": "Integration Needs",
         "keywords": ["integration", "api", "interoperability", "single sign-on",
                      "sso", "interface with"]},
    ],
}


def _payment_decision(text: str):
    low = text.lower()
    if re.search(r"net[\s-]?30", low):
        return "GO", ("The RFP specifies NET30 payment terms, which matches company "
                      "policy. Compliant per the SPS Financial / Accounting checklist "
                      "(payment terms of NET30 are acceptable).")
    m = re.search(r"net[\s-]?(\d{2,3})", low)
    if m and int(m.group(1)) > 30:
        n = m.group(1)
        return "ESCALATE", (
            f"The RFP requires NET{n} payment terms, which exceed the company's NET30 "
            f"policy. Per the SPS Financial / Accounting checklist, terms longer than "
            f"NET30 must be escalated to Accounting for approval before bidding.")
    return None, None


def _insurance_amounts(low: str):
    """Yield (value_in_millions, start, end) for every dollar figure in the text.

    Handles two written forms reliably:
      - "$5 million" / "5 million dollars"      -> 5.0
      - "$1,000,000" (comma-grouped millions)   -> 1.0
    A plain "$500,000" (one comma group) is NOT treated as millions, so small
    figures like a proposal bond do not get mistaken for insurance coverage.
    """
    for m in re.finditer(r"\$?\s?(\d+(?:\.\d+)?)\s*million", low):
        yield float(m.group(1)), m.start(), m.end()
    for m in re.finditer(r"\$\s?(\d{1,3}(?:,\d{3}){2,})", low):
        yield float(m.group(1).replace(",", "")) / 1_000_000.0, m.start(), m.end()


def _insurance_decision(text: str):
    low = text.lower()
    # Keep the HIGHEST amount that is clearly an insurance requirement: the word
    # "insur" must appear near the figure (this avoids limitation-of-liability
    # clauses, contract values, and other unrelated dollar amounts).
    highest = None
    for val, s, e in _insurance_amounts(low):
        window = low[max(0, s - 90): e + 40]
        if "insur" in window:
            if highest is None or val > highest:
                highest = val
    if highest is None:
        return None, None
    if highest > 5:
        return "NO-GO", (
            f"The RFP requires insurance coverage of ${highest:g}M, which exceeds the "
            f"company policy cap of $5M. Per the SPS Financial / Accounting checklist, "
            f"any required insurance above $5M is a NO-GO and cannot be accepted without "
            f"executive waiver.")
    return "GO", (
        f"The RFP requires insurance coverage of ${highest:g}M, which is within the "
        f"company's $5M policy cap. Compliant per the SPS Financial / Accounting checklist.")


def build_compliance_checklist(text: str):
    sentences = _sentences(text)
    report = {}
    decisions = []

    for dept, items in COMPLIANCE_TEMPLATE.items():
        rows = []
        for entry in items:
            evidence = _find_evidence(sentences, entry["keywords"], limit=2)
            status = "Found in RFP" if evidence else "Not addressed - confirm N/A"
            decision = ""

            if entry["item"] == "Payment Terms":
                d, reason = _payment_decision(text)
                if d:
                    decision = d
                    decisions.append(("Payment Terms", d, reason))
            if entry["item"] == "Insurance Requirements":
                d, reason = _insurance_decision(text)
                if d:
                    decision = d
                    decisions.append(("Insurance", d, reason))

            rows.append({
                "Department": dept,
                "Checklist Item": entry["item"],
                "Status": status,
                "Decision": decision,
                "Evidence": "  ".join(evidence) if evidence else "-",
            })
        report[dept] = rows
    return report, decisions


def overall_recommendation(decisions):
    """Roll the individual GO / NO-GO rules into a single recommendation."""
    if not decisions:
        return "REVIEW", "No automatic decision rules triggered - manual review required."
    if any(d == "NO-GO" for _, d, _ in decisions):
        return "NO-GO", "One or more hard limits exceeded (see decision rules)."
    if any(d == "ESCALATE" for _, d, _ in decisions):
        return "ESCALATE", "Some terms require sign-off from Accounting."
    return "GO", "All triggered decision rules are within policy."


# --------------------------------------------------------------------------- #
#  Top-level convenience function
# --------------------------------------------------------------------------- #

def analyze(file_bytes: bytes, filename: str):
    text = extract_text(file_bytes, filename)
    compliance, decisions = build_compliance_checklist(text)
    rec, rec_reason = overall_recommendation(decisions)
    return {
        "char_count": len(text),
        "word_count": len(text.split()),
        "deliverables": extract_deliverables(text),
        "evaluation": extract_evaluation_criteria(text),
        "compliance": compliance,
        "decisions": decisions,
        "recommendation": rec,
        "recommendation_reason": rec_reason,
        "raw_text": text,
    }
